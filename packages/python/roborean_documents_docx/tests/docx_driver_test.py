"""Tests for the DOCX document driver."""

import io

from docx import Document
from roborean_documents_docx.driver import (
    DocxDocumentDriver,
    iter_paragraphs,
)
from roborean_spec import DocumentOperation


def _replace_op(name: str, value: str) -> DocumentOperation:
    """Build a replace_named_value operation for tests."""
    return DocumentOperation.model_validate(
        {
            "documentId": "letter",
            "op": "replace_named_value",
            "name": name,
            "value": {
                "kind": "public_literal",
                "dataType": "string",
                "value": value,
            },
        }
    )


def _named_value_template_bytes() -> bytes:
    """Build a DOCX template with body, table, and header placeholders."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Dear ")
    paragraph.add_run("{{greeting}}")
    paragraph.add_run(", welcome.")

    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Note: {{table_note}}"

    section = document.sections[0]
    header = section.header
    header.paragraphs[0].text = "Title: {{header_title}}"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class TestDocxReplaceNamedValue:
    """Named-value substitution across the full paragraph tree."""

    def test_replaces_body_table_and_header(self) -> None:
        """Placeholders in body, table, and header are substituted."""
        driver = DocxDocumentDriver()
        driver._template = _named_value_template_bytes()

        session = driver.begin_session({}, {"documentId": "letter"})
        driver.apply_operation(session, _replace_op("greeting", "Ada"))
        driver.apply_operation(session, _replace_op("table_note", "On time"))
        driver.apply_operation(session, _replace_op("header_title", "Weekly"))
        driver.finalize(session)

        texts = [
            paragraph.text for paragraph in iter_paragraphs(session.document)
        ]
        assert "Dear Ada, welcome." in texts
        assert "Note: On time" in texts
        assert "Title: Weekly" in texts
        assert "{{greeting}}" not in "\n".join(texts)

    def test_preserves_formatting_outside_placeholder(self) -> None:
        """Runs outside the replaced span keep their formatting."""
        document = Document()
        paragraph = document.add_paragraph()
        bold = paragraph.add_run("Hello ")
        bold.bold = True
        paragraph.add_run("{{name}}")
        italic = paragraph.add_run("!")
        italic.italic = True

        buffer = io.BytesIO()
        document.save(buffer)

        driver = DocxDocumentDriver()
        driver._template = buffer.getvalue()
        session = driver.begin_session({}, {"documentId": "letter"})
        driver.apply_operation(session, _replace_op("name", "Ada"))
        driver.finalize(session)

        paragraph = session.document.paragraphs[0]
        assert paragraph.text == "Hello Ada!"
        assert paragraph.runs[0].bold is True
        assert paragraph.runs[0].text == "Hello "
        assert paragraph.runs[1].text == "Ada"
        assert paragraph.runs[2].italic is True
        assert paragraph.runs[2].text == "!"
