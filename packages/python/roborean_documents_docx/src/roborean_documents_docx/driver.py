"""DOCX document driver using docxtpl and python-docx."""

import io
import logging
from collections.abc import Iterator, Mapping
from dataclasses import dataclass, field
from datetime import UTC, datetime
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate
from pydantic import TypeAdapter
from roborean_documents_base.capabilities import assert_op_allowed
from roborean_documents_base.errors import DriverError
from roborean_documents_base.resolve_values import public_literal_value
from roborean_documents_base.template_store import DocumentTemplateStore
from roborean_spec import (
    DocumentDriverManifest,
    DocumentOperation,
    DocumentPreview,
    TemplateManifest,
    WorkspaceValue,
)

logger = logging.getLogger(__name__)


@dataclass
class DocxSession:
    """Working DOCX document session.

    Attributes:
        document_id: Document definition identifier for this session.
        driver_id: Driver id that owns the session.
        document: python-docx ``Document`` being edited.
        ops_applied: Serialized operations applied so far.
        context: Template context values rendered by docxtpl.
    """

    document_id: str
    driver_id: str
    document: Any
    ops_applied: list[dict[str, Any]] = field(default_factory=list)
    context: dict[str, Any] = field(default_factory=dict)


def iter_paragraphs(document: DocumentObject) -> Iterator[Paragraph]:
    """Yield paragraphs from body, tables, and section headers/footers.

    Args:
        document: Open python-docx document.

    Yields:
        Paragraph objects in body order, then table cells (including
        nested tables), then each section's non-linked header and footer.
    """

    def iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested in cell.tables:
                    yield from iter_table_paragraphs(nested)

    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        yield from iter_table_paragraphs(table)

    for section in document.sections:
        for header in (section.header, section.footer):
            if header.is_linked_to_previous:
                continue
            for paragraph in header.paragraphs:
                yield paragraph
            for table in header.tables:
                yield from iter_table_paragraphs(table)


def iter_block_items(parent: Any) -> Iterator[Paragraph | Table]:
    """Yield body paragraphs and tables in document order.

    Args:
        parent: ``Document`` or header/footer part exposing ``._element``.

    Yields:
        Top-level block items in XML order.
    """
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
    else:
        parent_element = parent._element

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _paragraph_html_tag(paragraph: Paragraph) -> str:
    """Map a Word paragraph style to an HTML tag name.

    Args:
        paragraph: Paragraph whose style name is inspected.

    Returns:
        ``h1``–``h6`` for heading styles, otherwise ``p``.
    """
    style = paragraph.style.name if paragraph.style else ""
    if style.startswith("Heading"):
        try:
            level = int(style.replace("Heading ", ""))
            return f"h{level}"
        except ValueError:
            return "h2"
    return "p"


def _replace_needle_in_paragraph(
    paragraph: Paragraph, needle: str, rendered: str
) -> bool:
    """Replace one placeholder span while preserving run formatting.

    Args:
        paragraph: Paragraph that may contain the placeholder needle.
        needle: Mustache-style placeholder such as ``{{name}}``.
        rendered: Public literal string to substitute.

    Returns:
        True when the needle was found and replaced.
    """
    full_text = paragraph.text
    start = full_text.find(needle)
    if start < 0:
        return False

    end = start + len(needle)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(full_text.replace(needle, rendered))
        return True

    run_starts: list[int] = []
    offset = 0
    for run in runs:
        run_starts.append(offset)
        offset += len(run.text)

    start_run = 0
    for index, run_start in enumerate(run_starts):
        run_end = run_start + len(runs[index].text)
        if run_start <= start < run_end:
            start_run = index
            break

    end_run = start_run
    for index, run_start in enumerate(run_starts):
        run_end = run_start + len(runs[index].text)
        if run_start < end <= run_end:
            end_run = index
            break

    start_offset = start - run_starts[start_run]
    end_offset = end - run_starts[end_run]

    if start_run == end_run:
        run = runs[start_run]
        run.text = run.text[:start_offset] + rendered + run.text[end_offset:]
        return True

    runs[start_run].text = runs[start_run].text[:start_offset] + rendered
    for index in range(start_run + 1, end_run):
        runs[index].text = ""
    runs[end_run].text = runs[end_run].text[end_offset:]
    return True


def _replace_named_value_in_document(
    document: DocumentObject, needle: str, rendered: str
) -> None:
    """Replace a named placeholder everywhere ``iter_paragraphs`` reaches.

    Args:
        document: Open python-docx document to mutate.
        needle: Mustache-style placeholder such as ``{{name}}``.
        rendered: Public literal string to substitute.
    """
    for paragraph in iter_paragraphs(document):
        if needle in paragraph.text:
            _replace_needle_in_paragraph(paragraph, needle, rendered)


def _render_table_html(table: Table) -> str:
    """Render one Word table as a minimal HTML table.

    Args:
        table: Table whose cell paragraphs are flattened to text.

    Returns:
        HTML ``table`` element string.
    """
    rows: list[str] = []
    for row in table.rows:
        cells = "".join(f"<td>{cell.text}</td>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    return f"<table>{''.join(rows)}</table>"


def _render_block_container_html(container: Any) -> str:
    """Render ordered paragraphs and tables from one document container.

    Args:
        container: ``Document`` body or a header/footer part.

    Returns:
        Concatenated HTML fragment for the container's block items.
    """
    parts: list[str] = []
    for block in iter_block_items(container):
        if isinstance(block, Paragraph):
            tag = _paragraph_html_tag(block)
            parts.append(f"<{tag}>{block.text}</{tag}>")
        else:
            parts.append(_render_table_html(block))
    return "".join(parts)


class DocxDocumentDriver:
    """Word document driver with flow ops and named slots.

    Attributes:
        driver_id: Stable driver identifier.
        manifest: Driver capability and media-type manifest.

        _template: Loaded ``.docx`` template bytes, if any.
        _manifest: Template sidecar manifest for required inputs.
    """

    driver_id: str = "roborean.docx"
    manifest: DocumentDriverManifest = DocumentDriverManifest(
        driverId="roborean.docx",
        version="0.3.0",
        irFamily="flow",
        capabilities=[
            "set_metadata",
            "replace_named_value",
            "flow.insert_paragraph",
            "flow.insert_heading",
            "finalize",
        ],
        supportsPreview=True,
        supportsBrowserExecution=False,
        supportsDiff=True,
        requiresBackend=True,
        templateMediaTypes=[
            "application/vnd.openxmlformats-officedocument"
            ".wordprocessingml.document"
        ],
    )

    _template: bytes | None
    _manifest: TemplateManifest | None

    def __init__(self) -> None:
        """Initialize empty template state."""
        self._template = None
        self._manifest = None

    def load_template(
        self,
        template_ref: str,
        *,
        store: DocumentTemplateStore,
        manifest: TemplateManifest,
    ) -> None:
        """Load ``.docx`` template bytes.

        Args:
            template_ref: Template identifier within the project package.
            store: Template store used to resolve template bytes.
            manifest: Validated template sidecar manifest.
        """
        self._template = store.load_bytes(template_ref)
        self._manifest = manifest

    def begin_session(
        self,
        workspace: Any,
        metadata: Mapping[str, Any],
    ) -> DocxSession:
        """Render docxtpl context then open a python-docx document.

        Args:
            workspace: Current workspace snapshot for template inputs.
            metadata: Session metadata such as ``documentId``.

        Returns:
            Open session bound to the rendered document.
        """
        assert self._template is not None

        # Open the template directly when there is nothing for docxtpl to bind.
        if self._manifest is None or not self._manifest.required_inputs:
            document = Document(io.BytesIO(self._template))
            return DocxSession(
                document_id=str(metadata.get("documentId", "")),
                driver_id=self.driver_id,
                document=document,
                context={},
            )

        # Render Jinja-style template slots from public workspace values.
        tpl = DocxTemplate(io.BytesIO(self._template))
        context: dict[str, Any] = {}
        if self._manifest is not None:
            for key in self._manifest.required_inputs:
                if key in workspace.values:
                    try:
                        context[key] = public_literal_value(
                            workspace.values[key]
                        )
                    except DriverError:
                        logger.debug(
                            "Skipping non-public template input %s",
                            key,
                            exc_info=True,
                        )
                        context[key] = ""

        # Materialize the rendered template into a python-docx Document.
        buffer = io.BytesIO()
        tpl.render(context)
        tpl.save(buffer)
        buffer.seek(0)
        document = Document(buffer)
        return DocxSession(
            document_id=str(metadata.get("documentId", "")),
            driver_id=self.driver_id,
            document=document,
            context=context,
        )

    def apply_operation(
        self, session: DocxSession, op: DocumentOperation
    ) -> None:
        """Apply flow / named-value operations.

        Args:
            session: Open DOCX session that receives the operation.
            op: Typed document operation to apply.

        Raises:
            UnsupportedOperationError: When the op is outside capabilities.
            DriverError: When a named-value payload is not a public literal.
        """
        assert_op_allowed(self.manifest, op)
        data = op.model_dump(mode="python", by_alias=True)
        session.ops_applied.append(op.model_dump(mode="json", by_alias=True))

        # Dispatch by operation name onto python-docx mutations.
        if op.op == "flow.insert_paragraph":
            text = "".join(run["text"] for run in data["runs"])
            session.document.add_paragraph(text)
        elif op.op == "flow.insert_heading":
            session.document.add_heading(
                str(data["text"]), level=int(data["level"])
            )
        elif op.op == "replace_named_value":
            value = TypeAdapter(WorkspaceValue).validate_python(data["value"])
            rendered = str(public_literal_value(value))
            needle = "{{" + str(data["name"]) + "}}"
            _replace_named_value_in_document(session.document, needle, rendered)
        elif op.op == "set_metadata":
            key = str(data["key"])
            value = TypeAdapter(WorkspaceValue).validate_python(data["value"])
            rendered = str(public_literal_value(value))
            props = session.document.core_properties
            if hasattr(props, key):
                setattr(props, key, rendered)

    def finalize(self, session: DocxSession) -> None:
        """No-op finalize.

        Args:
            session: Open session about to be serialized.
        """
        return

    def serialize(self, session: DocxSession) -> bytes:
        """Return ``.docx`` bytes.

        Args:
            session: Finalized session to serialize.

        Returns:
            Binary Office Open XML document payload.
        """
        buffer = io.BytesIO()
        session.document.save(buffer)
        return buffer.getvalue()

    def preview(self, session: DocxSession) -> DocumentPreview:
        """Approximate HTML from paragraphs, tables, and headers.

        Args:
            session: Finalized session to preview.

        Returns:
            HTML preview approximating paragraph and heading structure.
        """
        parts = ['<div class="roborean-docx">']
        parts.append(_render_block_container_html(session.document))

        for section in session.document.sections:
            for label, header in (
                ("header", section.header),
                ("footer", section.footer),
            ):
                if header.is_linked_to_previous:
                    continue
                fragment = _render_block_container_html(header)
                if fragment:
                    parts.append(
                        f'<div class="roborean-docx-{label}">{fragment}</div>'
                    )

        parts.append("</div>")

        return DocumentPreview(
            documentId=session.document_id,
            mode="html",
            body="".join(parts),
            warnings=[],
            generatedAt=datetime.now(UTC).isoformat(),
            renderer={
                "package": "roborean-documents-docx",
                "version": "0.3.0",
            },
        )


def create_driver() -> DocxDocumentDriver:
    """Entry-point factory.

    Returns:
        New ``DocxDocumentDriver`` instance.
    """
    return DocxDocumentDriver()


def docx_paragraph_texts(data: bytes) -> list[str]:
    """Extract paragraph texts for semantic compare.

    Args:
        data: Serialized ``.docx`` artifact bytes.

    Returns:
        Ordered list of paragraph text strings from body, tables, and
        section headers/footers.
    """
    document = Document(io.BytesIO(data))
    return [paragraph.text for paragraph in iter_paragraphs(document)]
